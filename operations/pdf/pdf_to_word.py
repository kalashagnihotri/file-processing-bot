import os
import logging

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

def convert_pdf_to_word(input_path, output_path):
    """
    Convert PDF to Word document using multiple methods with OCR fallback.
    """
    try:
        # Method 1: Try using pdf2docx (fastest for text-based PDFs)
        if PDF2DOCX_AVAILABLE:
            try:
                cv = Converter(input_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()
                return f"Converted PDF to Word using pdf2docx: {output_path}"
            except Exception as e:
                logging.warning(f"pdf2docx failed: {e}, trying OCR method")
        
        # Method 2: OCR-based conversion using Tesseract
        if PYMUPDF_AVAILABLE and PYTHON_DOCX_AVAILABLE and PYTESSERACT_AVAILABLE:
            try:
                from PIL import Image
                import io
                
                # Open the PDF
                pdf_document = fitz.open(input_path)
                doc = Document()
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    
                    # First try to extract text directly
                    text = page.get_text()
                    
                    if text.strip():
                        # If text exists, use it directly
                        doc.add_paragraph(text)
                    else:
                        # If no text, use OCR on the page image
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        image = Image.open(io.BytesIO(img_data))
                        
                        # Extract text using OCR
                        ocr_text = pytesseract.image_to_string(image)
                        doc.add_paragraph(ocr_text)
                    
                    # Add page break except for the last page
                    if page_num < len(pdf_document) - 1:
                        doc.add_page_break()
                
                pdf_document.close()
                doc.save(output_path)
                return f"Converted PDF to Word using OCR (Tesseract): {output_path}"
                
            except Exception as e:
                logging.warning(f"Tesseract OCR failed: {e}, trying EasyOCR")
        
        # Method 3: OCR-based conversion using EasyOCR
        if PYMUPDF_AVAILABLE and PYTHON_DOCX_AVAILABLE and EASYOCR_AVAILABLE:
            try:
                from PIL import Image
                import numpy as np
                import io
                
                # Initialize EasyOCR reader
                reader = easyocr.Reader(['en'])
                
                # Open the PDF
                pdf_document = fitz.open(input_path)
                doc = Document()
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    
                    # Try to extract text directly first
                    text = page.get_text()
                    
                    if text.strip():
                        doc.add_paragraph(text)
                    else:
                        # Use OCR on the page image
                        pix = page.get_pixmap()
                        img_data = pix.tobytes("png")
                        image = Image.open(io.BytesIO(img_data))
                        
                        # Convert PIL image to numpy array for EasyOCR
                        img_array = np.array(image)
                        
                        # Extract text using EasyOCR
                        results = reader.readtext(img_array)
                        ocr_text = ' '.join([result[1] for result in results])
                        doc.add_paragraph(ocr_text)
                    
                    # Add page break except for the last page
                    if page_num < len(pdf_document) - 1:
                        doc.add_page_break()
                
                pdf_document.close()
                doc.save(output_path)
                return f"Converted PDF to Word using OCR (EasyOCR): {output_path}"
                
            except Exception as e:
                logging.warning(f"EasyOCR failed: {e}")
        
        # If all methods fail, return error message
        missing_deps = []
        if not PDF2DOCX_AVAILABLE:
            missing_deps.append("pdf2docx")
        if not PYMUPDF_AVAILABLE:
            missing_deps.append("PyMuPDF")
        if not PYTHON_DOCX_AVAILABLE:
            missing_deps.append("python-docx")
        if not PYTESSERACT_AVAILABLE and not EASYOCR_AVAILABLE:
            missing_deps.append("pytesseract or easyocr")
        
        return f"Error: PDF to Word conversion not available. Missing dependencies: {', '.join(missing_deps)}"
        
    except Exception as e:
        return f"Error converting PDF to Word: {e}"
